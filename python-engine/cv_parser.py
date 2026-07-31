import sys
import json
import os
import re

SKILL_PATTERNS = {
    # Languages
    'python': r'\bpython\b',
    'java': r'\bjava\b',
    'javascript': r'\bjavascript\b',
    'c++': r'\bc\s*\+{2}\b',
    'c#': r'\bc\s*#\b',
    'typescript': r'\btypescript\b',
    'go': r'\bgo\b',
    'rust': r'\brust\b',
    'kotlin': r'\bkotlin\b',
    'swift': r'\bswift\b',
    'r': r'\br\b',
    'matlab': r'\bmatlab\b',
    'scala': r'\bscala\b',
    'perl': r'\bperl\b',
    'php': r'\bphp\b',
    'ruby': r'\bruby\b',
    'bash': r'\bbash\b',
    'shell': r'\bshell\b',
    'powershell': r'\bpowershell\b',
    'sql': r'\bsql\b',
    'html': r'\bhtml\b',
    'css': r'\bcss\b',
    'sass': r'\bsass\b',
    'less': r'\bless\b',
    'dart': r'\bdart\b',

    # Frontend frameworks/libraries
    'react': r'\breact\b',
    'angular': r'\bangular\b',
    'vue': r'\bvue\b',
    'svelte': r'\bsvelte\b',
    'jquery': r'\bjquery\b',
    'bootstrap': r'\bbootstrap\b',
    'tailwind': r'\btailwind\b',
    'redux': r'\bredux\b',
    'next.js': r'\bnext\s*\.?\s*js\b',
    'nuxt': r'\bnuxt\b',

    # Backend frameworks
    'node.js': r'\bnode\s*\.?\s*js\b',
    'express': r'\bexpress\b',
    'django': r'\bdjango\b',
    'flask': r'\bflask\b',
    'spring': r'\bspring\b',
    'spring boot': r'\bspring\s*boot\b',
    'asp.net': r'\basp\s*\.?\s*net\b',
    'laravel': r'\blaravel\b',
    'rails': r'\brails\b',
    'fastapi': r'\bfastapi\b',

    # Databases
    'mongodb': r'\bmongodb\b',
    'postgresql': r'\bpostgresql\b',
    'mysql': r'\bmysql\b',
    'sqlite': r'\bsqlite\b',
    'oracle': r'\boracle\b',
    'microsoft sql server': r'\bmicrosoft\s*sql\s*server\b',
    'mariadb': r'\bmariadb\b',
    'redis': r'\bredis\b',
    'cassandra': r'\bcassandra\b',
    'dynamodb': r'\bdynamodb\b',
    'firebase': r'\bfirebase\b',
    'supabase': r'\bsupabase\b',
    'neo4j': r'\bneo4j\b',
    'couchdb': r'\bcouchdb\b',

    # Cloud & DevOps
    'aws': r'\baws\b',
    'azure': r'\bazure\b',
    'gcp': r'\bgcp\b',
    'google cloud': r'\bgoogle\s*cloud\b',
    'docker': r'\bdocker\b',
    'kubernetes': r'\bkubernetes\b',
    'terraform': r'\bterraform\b',
    'ansible': r'\bansible\b',
    'jenkins': r'\bjenkins\b',
    'github actions': r'\bgithub\s*actions\b',
    'circleci': r'\bcircleci\b',
    'gitlab ci': r'\bgitlab\s*ci\b',
    'docker compose': r'\bdocker\s*compose\b',
    'helm': r'\bhelm\b',
    'prometheus': r'\bprometheus\b',
    'grafana': r'\bgrafana\b',
    'git': r'\bgit\b',
    'linux': r'\blinux\b',
    'nginx': r'\bnginx\b',

    # Data Science & ML
    'machine learning': r'\bmachine\s*learning\b',
    'deep learning': r'\bdeep\s*learning\b',
    'data science': r'\bdata\s*science\b',
    'data analysis': r'\bdata\s*analysis\b',
    'data engineering': r'\bdata\s*engineering\b',
    'data visualization': r'\bdata\s*visualization\b',
    'statistics': r'\bstatistics\b',
    'nlp': r'\bnlp\b',
    'computer vision': r'\bcomputer\s*vision\b',
    'tensorflow': r'\btensorflow\b',
    'pytorch': r'\bpytorch\b',
    'keras': r'\bkeras\b',
    'scikit-learn': r'\bscikit\s*[- ]?\s*learn\b',
    'pandas': r'\bpandas\b',
    'numpy': r'\bnumpy\b',
    'matplotlib': r'\bmatplotlib\b',
    'seaborn': r'\bseaborn\b',
    'plotly': r'\bplotly\b',
    'opencv': r'\bopencv\b',
    'llm': r'\bllm\b',
    'langchain': r'\blangchain\b',
    'tableau': r'\btableau\b',
    'power bi': r'\bpower\s*bi\b',
    'excel': r'\bexcel\b',
    'spark': r'\bspark\b',
    'hadoop': r'\bhadoop\b',
    'airflow': r'\bairflow\b',
    'etl': r'\betl\b',

    # Testing
    'jest': r'\bjest\b',
    'mocha': r'\bmocha\b',
    'cypress': r'\bcypress\b',
    'selenium': r'\bselenium\b',
    'playwright': r'\bplaywright\b',
    'junit': r'\bjunit\b',
    'pytest': r'\bpytest\b',

    # Tools & Platforms
    'rest api': r'\brestful?\s*apis?\b',
    'graphql': r'\bgraphql\b',
    'grpc': r'\bgrpc\b',
    'websocket': r'\bwebsocket\b',
    'jira': r'\bjira\b',
    'confluence': r'\bconfluence\b',
    'postman': r'\bpostman\b',
    'swagger': r'\bswagger\b',
    'figma': r'\bfigma\b',
    'sketch': r'\bsketch\b',
    'photoshop': r'\bphotoshop\b',
    'docker hub': r'\bdocker\s*hub\b',
    'npm': r'\bnpm\b',
    'webpack': r'\bwebpack\b',
    'vite': r'\bvite\b',
    'babel': r'\bbabel\b',
    'eslint': r'\beslint\b',
    'prettier': r'\bprettier\b',

    # CI/CD & Agile
    'ci/cd': r'\bci\s*/\s*cd\b',
    'agile': r'\bagile\b',
    'scrum': r'\bscrum\b',
    'kanban': r'\bkanban\b',
    'microservices': r'\bmicroservices?\b',
    'serverless': r'\bserverless\b',

    # Soft skills
    'communication': r'\bcommunication\b',
    'leadership': r'\bleadership\b',
    'teamwork': r'\bteamwork\b',
    'problem solving': r'\bproblem\s*solving\b',
    'critical thinking': r'\bcritical\s*thinking\b',
    'project management': r'\bproject\s*management\b',
    'time management': r'\btime\s*management\b',
}

def extract_skills(text):
    text_clean = re.sub(r'[^\w\s#+]', ' ', text.lower())
    text_clean = re.sub(r'\s+', ' ', text_clean)
    found = []
    for skill_name, pattern in SKILL_PATTERNS.items():
        if re.search(pattern, text_clean):
            found.append(skill_name)
    return sorted(set(found))


def extract_text_from_pdf(filepath):
    import fitz
    doc = fitz.open(filepath)
    text = ''
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def _extract_docx_table_text(table):
    text = ''
    for row in table.rows:
        for cell in row.cells:
            text += cell.text + ' '
    return text


def extract_text_from_docx(filepath):
    import docx
    doc = docx.Document(filepath)
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        parts.append(_extract_docx_table_text(table))
    return '\n'.join(parts)


if __name__ == '__main__':
    filepath = ''
    for i in range(len(sys.argv)):
        if sys.argv[i] == '--file' and i + 1 < len(sys.argv):
            filepath = sys.argv[i + 1]

    if not filepath or not os.path.exists(filepath):
        print(json.dumps({'error': 'File not found', 'file': filepath}))
        sys.exit(1)

    ext = os.path.splitext(filepath)[1].lower()
    text = ''

    if ext == '.pdf':
        text = extract_text_from_pdf(filepath)
    elif ext in ('.docx', '.doc'):
        text = extract_text_from_docx(filepath)
    else:
        print(json.dumps({'error': f'Unsupported file type: {ext}'}))
        sys.exit(1)

    skills = extract_skills(text)
    print(json.dumps({'skills': skills, 'textLength': len(text)}))
